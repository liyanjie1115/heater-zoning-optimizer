from __future__ import annotations

import os
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PySide6.QtWidgets import QApplication, QTabWidget
from matplotlib import patches

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from src.heater_zoning.config import AnalysisConfig
from src.heater_zoning.exporters import export_summary_pdf
from src.heater_zoning.fonts import configure_matplotlib_fonts
from src.heater_zoning.gui import APP_STYLESHEET, DesktopApp
from src.heater_zoning.runflow import run_analysis_pipeline
from src.heater_zoning.visualization import (
    build_metrics_bar_matplotlib,
    build_metrics_radar_matplotlib,
    build_module_layout_matplotlib,
    build_temperature_comparison_matplotlib,
)


DOC_ROOT = PROJECT_ROOT / "docs" / "softcopyright_manual"
ASSET_DIR = DOC_ROOT / "assets"
OUTPUT_DIR = DOC_ROOT / "generated"
DOCX_PATH = DOC_ROOT / "Heater_Zoning_Optimizer_V1.0_软件使用说明书.docx"
PDF_PATH = DOC_ROOT / "Heater_Zoning_Optimizer_V1.0_软件使用说明书.pdf"


def ensure_dirs():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def save_figure(fig, path: Path, dpi: int = 160):
    fig.savefig(path, dpi=dpi, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def build_table_figure(dataframe: pd.DataFrame, title: str, path: Path):
    rows = min(len(dataframe), 12)
    fig, axis = plt.subplots(figsize=(11.2, max(3.2, rows * 0.42 + 1.3)))
    fig.patch.set_facecolor("white")
    axis.axis("off")
    axis.set_title(title, loc="left", fontsize=14, fontweight="bold", pad=12)
    table = axis.table(
        cellText=dataframe.head(rows).values,
        colLabels=[str(column) for column in dataframe.columns],
        loc="center",
        cellLoc="center",
    )
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 1.45)
    for (row, col), cell in table.get_celld().items():
        if row == 0:
            cell.set_facecolor("#d9eaf7")
            cell.set_text_props(weight="bold")
        cell.set_edgecolor("#cbd5e1")
    fig.tight_layout()
    save_figure(fig, path)


def build_module_diagram(path: Path):
    fig, ax = plt.subplots(figsize=(11.2, 6.4))
    fig.patch.set_facecolor("white")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")

    modules = [
        (0.6, 7.2, 2.4, 1.2, "#dbeafe", "输入与参数层"),
        (3.2, 7.2, 2.4, 1.2, "#dcfce7", "分析计算层"),
        (5.8, 7.2, 2.4, 1.2, "#fef3c7", "结果呈现层"),
        (8.4, 7.2, 1.0, 1.2, "#fee2e2", "导出层"),
        (0.8, 5.0, 2.0, 0.95, "#eff6ff", "示例数据/本地数据"),
        (3.4, 5.0, 2.0, 0.95, "#ecfccb", "分区方案生成"),
        (5.9, 5.0, 2.0, 0.95, "#fef9c3", "评分与差异分析"),
        (8.15, 5.0, 1.4, 0.95, "#ffe4e6", "Excel/PDF"),
        (0.8, 3.2, 2.0, 0.95, "#eff6ff", "参数模板/最近文件"),
        (3.4, 3.2, 2.0, 0.95, "#ecfccb", "等距分区"),
        (5.9, 3.2, 2.0, 0.95, "#fef9c3", "模块对齐分区"),
        (8.15, 3.2, 1.4, 0.95, "#ffe4e6", "图表与明细"),
    ]

    for x, y, w, h, color, text in modules:
        rect = patches.FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.02,rounding_size=0.08",
            facecolor=color,
            edgecolor="#64748b",
            linewidth=1.1,
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=11, fontweight="bold")

    arrows = [
        ((2.8, 5.47), (3.4, 5.47)),
        ((5.4, 5.47), (5.9, 5.47)),
        ((7.9, 5.47), (8.15, 5.47)),
        ((2.8, 3.67), (3.4, 3.67)),
        ((5.4, 3.67), (5.9, 3.67)),
        ((8.15, 4.15), (7.2, 6.95)),
        ((4.4, 4.95), (4.4, 4.15)),
        ((6.9, 4.95), (6.9, 4.15)),
    ]
    for start, end in arrows:
        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", lw=1.5, color="#475569"))

    ax.text(0.4, 9.2, "Heater Zoning Optimizer 功能模块图", fontsize=17, fontweight="bold", ha="left")
    ax.text(0.4, 8.7, "展示软件从输入、分析、展示到导出的主要功能结构。", fontsize=11, color="#475569", ha="left")
    fig.tight_layout()
    save_figure(fig, path)


def build_flow_diagram(path: Path):
    fig, ax = plt.subplots(figsize=(8.8, 10.8))
    fig.patch.set_facecolor("white")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 14)
    ax.axis("off")

    steps = [
        (2.1, 12.2, 5.8, 0.95, "#dbeafe", "启动桌面程序"),
        (2.1, 10.6, 5.8, 0.95, "#e0f2fe", "选择示例数据或导入本地文件"),
        (2.1, 9.0, 5.8, 0.95, "#ecfccb", "配置参数模板与分析参数"),
        (2.1, 7.4, 5.8, 0.95, "#fef3c7", "执行分区计算与评分分析"),
        (2.1, 5.8, 5.8, 0.95, "#fde68a", "查看推荐方案、图表与工程明细"),
        (2.1, 4.2, 5.8, 0.95, "#fecaca", "导出 Excel 报告与 PDF 摘要"),
        (2.1, 2.6, 5.8, 0.95, "#ddd6fe", "结束/归档分析结果"),
    ]

    for x, y, w, h, color, text in steps:
        rect = patches.FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.03,rounding_size=0.08",
            facecolor=color,
            edgecolor="#64748b",
            linewidth=1.1,
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=12, fontweight="bold")

    for idx in range(len(steps) - 1):
        x = 5.0
        y1 = steps[idx][1]
        y2 = steps[idx + 1][1] + steps[idx + 1][3]
        ax.annotate("", xy=(x, y2), xytext=(x, y1), arrowprops=dict(arrowstyle="->", lw=1.6, color="#475569"))

    ax.text(0.6, 13.3, "Heater Zoning Optimizer 业务操作流程图", fontsize=17, fontweight="bold", ha="left")
    ax.text(0.6, 12.8, "描述用户从启动、导入、分析到导出的完整主流程。", fontsize=11, color="#475569", ha="left")
    fig.tight_layout()
    save_figure(fig, path)


def create_static_assets(artifacts):
    configure_matplotlib_fonts()
    result = artifacts.result

    build_module_diagram(ASSET_DIR / "functional_module_diagram.png")
    build_flow_diagram(ASSET_DIR / "operation_flow_diagram.png")
    save_figure(
        build_temperature_comparison_matplotlib(result.profile_df, result.equal_zones, result.aligned_zones),
        ASSET_DIR / "temperature_comparison.png",
    )
    save_figure(
        build_module_layout_matplotlib(result.equal_zones, result.aligned_zones),
        ASSET_DIR / "module_layout.png",
    )
    save_figure(
        build_metrics_bar_matplotlib(result.equal_metrics, result.aligned_metrics),
        ASSET_DIR / "metrics_bar.png",
    )
    save_figure(
        build_metrics_radar_matplotlib(result.equal_metrics, result.aligned_metrics),
        ASSET_DIR / "metrics_radar.png",
    )

    build_table_figure(result.profile_df.head(15), "示例温度剖面数据（前 15 行）", ASSET_DIR / "sample_profile_table.png")
    build_table_figure(artifacts.frames.metrics, "评价指标明细", ASSET_DIR / "metrics_table.png")
    build_table_figure(artifacts.frames.differences, "方案差异摘要", ASSET_DIR / "differences_table.png")
    build_table_figure(artifacts.zone_summary, "分区结果汇总", ASSET_DIR / "zone_summary_table.png")


def _find_tab_widget(window: DesktopApp, expected_titles: list[str]) -> QTabWidget:
    for widget in window.findChildren(QTabWidget):
        titles = [widget.tabText(index) for index in range(widget.count())]
        if titles == expected_titles:
            return widget
    raise RuntimeError(f"Tab widget not found: {expected_titles}")


def capture_window(window: DesktopApp, path: Path):
    QApplication.processEvents()
    window.repaint()
    QApplication.processEvents()
    window.grab().save(str(path))


def create_ui_screenshots(artifacts):
    app = QApplication.instance() or QApplication([])
    app.setStyleSheet(APP_STYLESHEET)

    window = DesktopApp()
    window.current_artifacts = artifacts
    window.recommended_label.setText(artifacts.summary_cards[0]["value"])
    window.excel_path_label.setText(str(artifacts.export_path))
    window.status_label.setText("分析完成")
    window.status_hint_label.setText("已基于默认示例数据生成结果，可查看图表、明细与导出文件。")
    window._update_summary_cards(artifacts)
    window._update_decision_cards(artifacts)
    window._update_tables(artifacts)
    window._update_charts(artifacts)

    window.show()
    QApplication.processEvents()

    main_tabs = _find_tab_widget(window, ["决策总览", "图表", "工程明细"])
    engineering_tabs = _find_tab_widget(window, ["分区汇总", "分区结果", "采样点"])

    main_tabs.setCurrentIndex(0)
    capture_window(window, ASSET_DIR / "ui_decision_overview.png")

    main_tabs.setCurrentIndex(1)
    capture_window(window, ASSET_DIR / "ui_charts.png")

    main_tabs.setCurrentIndex(2)
    engineering_tabs.setCurrentIndex(0)
    capture_window(window, ASSET_DIR / "ui_engineering_summary.png")

    engineering_tabs.setCurrentIndex(1)
    capture_window(window, ASSET_DIR / "ui_engineering_zones.png")

    window.close()
    QApplication.processEvents()


def set_document_style(document: Document):
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        heading_style = document.styles[style_name]
        heading_style.font.name = "黑体"
        heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_title(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_center_paragraph(document: Document, text: str, size: int = 12, bold: bool = False, font_name: str = "宋体"):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_caption(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_image(document: Document, path: Path, width_cm: float, caption: str):
    document.add_picture(str(path), width=Cm(width_cm))
    add_caption(document, caption)


def add_key_value_table(document: Document, rows: list[tuple[str, str]]):
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "项目"
    header[1].text = "说明"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value


def add_dataframe_table(document: Document, dataframe: pd.DataFrame, max_rows: int = 8):
    preview = dataframe.head(max_rows).fillna("")
    table = document.add_table(rows=1, cols=len(preview.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, column in enumerate(preview.columns):
        table.rows[0].cells[index].text = str(column)
    for _, row in preview.iterrows():
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = f"{value:.4f}" if isinstance(value, float) else str(value)


def generate_document(artifacts):
    result = artifacts.result
    document = Document()
    set_document_style(document)

    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    document.add_paragraph("")
    document.add_paragraph("")
    add_center_paragraph(document, "计算机软件著作权登记", size=18, bold=True, font_name="黑体")
    document.add_paragraph("")
    add_center_paragraph(document, "Heater Zoning Optimizer V1.0", size=22, bold=True, font_name="黑体")
    add_center_paragraph(document, "软件使用说明书", size=22, bold=True, font_name="黑体")
    document.add_paragraph("")
    document.add_paragraph("")
    add_center_paragraph(document, "软件名称：Heater Zoning Optimizer", size=14)
    add_center_paragraph(document, "版本号：V1.0.0", size=14)
    add_center_paragraph(document, "文档性质：软件著作权登记申请材料", size=14)
    add_center_paragraph(document, "完成日期：2026年5月8日", size=14)
    document.add_page_break()

    document.add_heading("版权页", level=1)
    document.add_paragraph("软件名称：Heater Zoning Optimizer")
    document.add_paragraph("版本号：V1.0.0")
    document.add_paragraph("文档名称：Heater Zoning Optimizer V1.0 软件使用说明书")
    document.add_paragraph("著作权说明：本文档及其所附图示、界面截图、样例输出均用于软件著作权登记与软件交付说明。")
    document.add_paragraph("文档用途：说明软件运行环境、主要功能、操作方式及输出结果。")
    document.add_paragraph("保留权利：未经许可，不得擅自复制、传播或用于与本软件无关之用途。")
    document.add_page_break()

    document.add_heading("1. 编写目的", level=1)
    document.add_paragraph(
        "本文档用于说明 Heater Zoning Optimizer V1.0 的安装方式、运行环境、主要功能、操作流程、输入输出要求及注意事项，"
        "供软件著作权登记、内部交付和最终用户培训使用。"
    )

    document.add_heading("2. 软件概述", level=1)
    document.add_paragraph(
        "本软件用于分析加热区域沿长度方向的温度剖面，在模块长度、模块间距、边缘外伸余量等安装约束下，"
        "对等距分区方案与模块对齐分区方案进行对比，给出推荐方案、量化评分、图表展示与报告导出结果。"
    )
    add_key_value_table(
        document,
        [
            ("软件定位", "加热分区优化分析桌面软件"),
            ("主要形态", "Windows 桌面 GUI，支持 CLI 和可选 Web 视图"),
            ("核心输出", "推荐方案、分区结果、评价指标、差异摘要、Excel 报告、PDF 摘要"),
            ("当前示例结论", artifacts.summary_cards[0]["value"]),
        ],
    )

    document.add_heading("3. 术语与缩略语", level=1)
    add_key_value_table(
        document,
        [
            ("温度剖面", "沿加热区域长度方向采集得到的距离-温度序列数据"),
            ("等距分区", "按固定区间长度将整体加热区域进行等间距划分"),
            ("模块对齐分区", "结合模块长度、模块间距和边缘余量约束进行对齐划分"),
            ("评价指标", "用于比较两类分区方案优劣的量化评分项"),
            ("工程明细", "包含分区边界、采样点、模块数量及误差等工程核查信息"),
        ],
    )

    document.add_heading("4. 运行环境", level=1)
    add_key_value_table(
        document,
        [
            ("操作系统", "Windows 10 / Windows 11"),
            ("运行方式", "安装版 setup.exe 或免安装版 heater-zoning-optimizer.exe"),
            ("建议分辨率", "1920 x 1080 及以上"),
            ("Python 环境", "开发环境使用 Python 3.11（文档生成基于项目当前虚拟环境）"),
            ("主要依赖", "PySide6、pandas、matplotlib、openpyxl、Flask、plotly"),
        ],
    )

    document.add_heading("5. 功能模块设计", level=1)
    document.add_paragraph(
        "软件采用输入与参数层、分析计算层、结果呈现层和导出层四级结构。输入与参数层负责数据与模板管理；"
        "分析计算层负责生成等距分区和模块对齐分区并完成评分；结果呈现层负责展示总览、图表和工程明细；"
        "导出层负责输出 Excel 报告和 PDF 摘要。"
    )
    add_image(document, ASSET_DIR / "functional_module_diagram.png", 16.2, "图 1 软件功能模块图")

    document.add_heading("6. 安装与启动", level=1)
    document.add_paragraph("4.1 安装版使用方法：")
    document.add_paragraph(
        "双击 release/installer 目录下生成的安装包，按照安装向导完成安装。安装完成后，可通过桌面快捷方式或开始菜单启动软件。",
        style=None,
    )
    document.add_paragraph("4.2 免安装版使用方法：")
    document.add_paragraph("双击 heater-zoning-optimizer.exe 即可直接启动。")
    document.add_paragraph("4.3 首次启动建议：")
    document.add_paragraph("首次运行时，建议先选择默认示例数据执行一次分析，以确认界面、图表和导出链路正常。")

    document.add_heading("7. 输入数据要求", level=1)
    document.add_paragraph("软件支持 CSV、XLSX、XLS 格式温度剖面文件，至少包含以下两列：")
    document.add_paragraph("1. distance_mm：测点距离，单位 mm")
    document.add_paragraph("2. temperature_c：测点温度，单位 ℃")
    document.add_paragraph("同时兼容 distance、x、距离、temperature、temp、温度 等别名列名。")
    add_image(document, ASSET_DIR / "sample_profile_table.png", 15.8, "图 2 示例温度剖面数据预览")

    document.add_heading("8. 软件界面与功能结构", level=1)
    document.add_paragraph(
        "软件主界面由左侧参数控制区和右侧结果展示区组成。左侧用于选择模板、输入数据、配置参数与执行分析；"
        "右侧用于展示推荐结果、图表、工程明细与导出状态。"
    )
    add_image(document, ASSET_DIR / "ui_decision_overview.png", 16.2, "图 3 软件主界面与决策总览")

    document.add_heading("9. 操作流程说明", level=1)
    document.add_paragraph("图 4 所示为软件标准使用流程。")
    add_image(document, ASSET_DIR / "operation_flow_diagram.png", 14.0, "图 4 软件操作流程图")
    document.add_paragraph("7.1 选择参数模板")
    document.add_paragraph(
        "软件内置平衡推荐、追求贴合、优先安装、控制复杂度等参数模板。选择模板后，分析参数会自动写入当前表单。"
    )
    document.add_paragraph("7.2 选择数据来源")
    document.add_paragraph("用户可选择“使用示例数据”，也可选择本地 CSV/Excel 文件作为分析输入。")
    document.add_paragraph("7.3 设置分析参数")
    document.add_paragraph(
        "基础参数页签包括总长度、最大分区数、等距分区数、模块长度、模块间距等；专家参数页签包括拟合权重、分离权重、梯度权重等高级项。"
    )
    document.add_paragraph("7.4 运行分析")
    document.add_paragraph("点击“运行分析”按钮后，软件会计算分区结果、评分指标、图表和导出文件。")

    document.add_paragraph("7.5 查看决策总览")
    document.add_paragraph(
        "决策总览页集中展示推荐方案、推荐把握度、两类方案得分、当前分析设置、评价指标明细和方案差异摘要。"
    )
    add_image(document, ASSET_DIR / "ui_decision_overview.png", 16.2, "图 5 决策总览页")

    document.add_paragraph("7.6 查看图表")
    document.add_paragraph(
        "图表页展示温度与分区对比图、模块布局图、归一化指标对比图和质量雷达图，用于快速理解两类方案的差别。"
    )
    add_image(document, ASSET_DIR / "ui_charts.png", 16.2, "图 6 图表页")
    add_image(document, ASSET_DIR / "temperature_comparison.png", 15.6, "图 7 温度与分区对比图")
    add_image(document, ASSET_DIR / "module_layout.png", 15.6, "图 8 模块布局图")

    document.add_paragraph("7.7 查看工程明细")
    document.add_paragraph(
        "工程明细页包括分区汇总、分区结果和采样点明细，可用于工程人员核查边界位置、采样点布置和模块匹配情况。"
    )
    add_image(document, ASSET_DIR / "ui_engineering_summary.png", 16.2, "图 9 工程明细-分区汇总")
    add_image(document, ASSET_DIR / "ui_engineering_zones.png", 16.2, "图 10 工程明细-分区结果")

    document.add_heading("10. 示例结果说明", level=1)
    document.add_paragraph("基于软件默认示例数据，系统自动生成如下分析结论：")
    add_key_value_table(
        document,
        [
            ("推荐方案", artifacts.summary_cards[0]["value"]),
            ("推荐把握度", artifacts.summary_cards[1]["value"]),
            ("等距分区综合得分", artifacts.summary_cards[2]["value"]),
            ("模块对齐综合得分", artifacts.summary_cards[3]["value"]),
            ("综合得分差值", artifacts.summary_cards[4]["value"]),
            ("Excel 报告输出", str(artifacts.export_path)),
        ],
    )
    document.add_paragraph("")
    add_image(document, ASSET_DIR / "metrics_bar.png", 15.6, "图 11 归一化指标对比图")
    add_image(document, ASSET_DIR / "metrics_radar.png", 15.6, "图 12 质量雷达图")
    add_image(document, ASSET_DIR / "zone_summary_table.png", 15.6, "图 13 分区结果汇总表")
    add_image(document, ASSET_DIR / "differences_table.png", 15.6, "图 14 方案差异摘要表")

    document.add_heading("11. 导出文件说明", level=1)
    document.add_paragraph("软件在完成分析后会自动生成 Excel 报告，并支持导出 PDF 摘要。")
    add_key_value_table(
        document,
        [
            ("Excel 报告路径", str(artifacts.export_path)),
            ("PDF 摘要路径", str(OUTPUT_DIR / "manual_sample_summary.pdf")),
            ("Excel 主要工作表", "原始数据、等距分区结果、模块对齐分区结果、三点采样、评价指标、方案差异摘要"),
            ("PDF 主要内容", "摘要页、分区对比图、模块布局图、评价图表及关键表格"),
        ],
    )

    document.add_heading("12. 异常处理与注意事项", level=1)
    document.add_paragraph("1. 若输入文件缺少 distance_mm 或 temperature_c 列，软件将提示数据格式错误。")
    document.add_paragraph("2. 若首次运行出现 Windows 安全提示，可在确认来源可信后继续运行。")
    document.add_paragraph("3. 若导出失败，应检查输出目录写权限及目标文件是否被其他程序占用。")
    document.add_paragraph("4. 若需批量自动执行分析，可使用项目提供的 CLI 入口。")

    document.add_heading("13. 附录：示例数据与指标预览", level=1)
    document.add_paragraph("附表 1 为示例数据前 8 行，附表 2 为评价指标前 8 行。")
    add_dataframe_table(document, result.profile_df[["distance_mm", "temperature_c"]], max_rows=8)
    document.add_paragraph("")
    add_dataframe_table(document, artifacts.frames.metrics, max_rows=8)

    document.save(DOCX_PATH)


def export_docx_to_pdf(docx_path: Path, pdf_path: Path):
    import win32com.client  # type: ignore

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    document = None
    try:
        document = word.Documents.Open(str(docx_path.resolve()))
        document.ExportAsFixedFormat(
            OutputFileName=str(pdf_path.resolve()),
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,
            Range=0,
            Item=0,
            IncludeDocProps=True,
            KeepIRM=True,
            CreateBookmarks=1,
            DocStructureTags=True,
            BitmapMissingFonts=True,
            UseISO19005_1=False,
        )
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()


def main():
    ensure_dirs()
    config = AnalysisConfig().validate()
    artifacts = run_analysis_pipeline(
        config=config,
        source="sample",
        output_dir=OUTPUT_DIR,
        output_name="manual_sample_report.xlsx",
    )
    export_summary_pdf(artifacts.result, OUTPUT_DIR / "manual_sample_summary.pdf")
    create_static_assets(artifacts)
    create_ui_screenshots(artifacts)
    generate_document(artifacts)
    export_docx_to_pdf(DOCX_PATH, PDF_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    os.chdir(PROJECT_ROOT)
    main()
